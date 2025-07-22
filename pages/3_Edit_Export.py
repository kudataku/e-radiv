# ========== pages/3_Edit_Export.py ===================
import streamlit as st
import sqlite3
from docx import Document
from io import BytesIO
from utils.auth import require_login

require_login()
if st.session_state.role not in ["notulis", "admin"]:
    st.error("Anda tidak memiliki akses ke halaman ini.")
    st.stop()

st.title("✍️ Edit & Export Risalah")

conn = sqlite3.connect('database/risalah.db')
cursor = conn.cursor()
cursor.execute("SELECT id, tanggal, agenda, pimpinan FROM risalah ORDER BY id DESC")
rows = cursor.fetchall()

selected = st.selectbox("📂 Pilih risalah untuk diedit", rows, format_func=lambda x: f"{x[1]} - {x[2]} (Pimpinan: {x[3]})")

if selected:
    risalah_id = selected[0]
    cursor.execute("SELECT transcript FROM risalah WHERE id = ?", (risalah_id,))
    data = cursor.fetchone()
    edited_text = st.text_area("📝 Edit Transkrip", value=data[0], height=400)

    if st.button("💾 Simpan Perubahan"):
        cursor.execute("UPDATE risalah SET transcript = ? WHERE id = ?", (edited_text, risalah_id))
        conn.commit()
        st.success("Perubahan berhasil disimpan!")

    if st.button("📤 Export ke Word (.docx)"):
        tanggal, agenda, pimpinan = selected[1], selected[2], selected[3]
        doc = Document()
        doc.add_heading("RISALAH RAPAT DPRD KABUPATEN LEBAK", 0)
        doc.add_paragraph(f"Hari/Tanggal : {tanggal}")
        doc.add_paragraph(f"Agenda       : {agenda}")
        doc.add_paragraph(f"Pimpinan     : {pimpinan}")
        doc.add_paragraph("")
        doc.add_heading("Hasil Pembahasan:", level=1)
        doc.add_paragraph(edited_text)
        doc.add_paragraph("\n\nNotulis: _______________________")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="⬇️ Download Word",
            data=buffer,
            file_name=f"risalah_{tanggal.replace('-', '')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

conn.close()
